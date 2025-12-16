"""
Async document processing service using pgvector.

Replaces FAISS-based DocumentProcessor with PostgreSQL storage.
"""

import asyncio
import time
import uuid
from pathlib import Path
from typing import TYPE_CHECKING, Any, Callable, Optional
from urllib.parse import urlparse

from sqlalchemy import delete, select
from sqlalchemy.ext.asyncio import AsyncSession

from packages.core.config.settings import Config
from packages.core.database.models import Document, DocumentChunk, DocumentStatus
from apps.writer.services.documents.exceptions import (
    DocumentNotFoundError,
    FileOperationError,
    InvalidFilenameError,
    InvalidURLError,
    UnsupportedFileTypeError,
)
from packages.core.security.sanitization import is_safe_url, sanitize_filename
from packages.core.utils.async_io import delete_file_async, write_file_async
from apps.writer.services.vectorstore.pgvector_store import PgVectorStore

if TYPE_CHECKING:
    from packages.core.database.models import User


class DocumentService:
    """
    Async document processing service.

    Handles file parsing, chunking, embedding generation, and storage to pgvector.
    """

    ALLOWED_EXTENSIONS = {".pdf", ".txt", ".csv", ".md", ".docx", ".xlsx", ".png", ".jpg", ".jpeg"}

    def __init__(self, session: AsyncSession, config: Optional[Config] = None):
        self.session = session
        self.config = config or Config()

        # Lazy-loaded components
        self._text_splitter = None
        self._file_loaders = None

    # CRUD Operations

    async def get_document(self, document_id: uuid.UUID, user_id: uuid.UUID) -> Document:
        """Get a document by ID for a specific user."""
        stmt = select(Document).where(
            Document.id == document_id,
            Document.user_id == user_id,
        )
        result = await self.session.execute(stmt)
        document = result.scalar_one_or_none()

        if not document:
            raise DocumentNotFoundError(str(document_id))

        return document

    async def list_documents(self, user_id: uuid.UUID) -> list[Document]:
        """List all documents for a user."""
        stmt = (
            select(Document)
            .where(Document.user_id == user_id)
            .order_by(Document.created_at.desc())
        )
        result = await self.session.execute(stmt)
        return list(result.scalars().all())

    async def delete_document(
        self,
        document_id: uuid.UUID,
        user_id: uuid.UUID,
    ) -> None:
        """Delete a document and its file."""
        document = await self.get_document(document_id, user_id)

        try:
            if document.storage_key and not document.storage_key.startswith("http"):
                file_path = Path(document.storage_key)
                if file_path.exists():
                    await delete_file_async(file_path)

            await self.session.delete(document)
            await self.session.commit()

        except DocumentNotFoundError:
            raise
        except Exception as e:
            await self.session.rollback()
            raise FileOperationError(f"Failed to delete document: {e}")

    async def delete_all_documents(self, upload_dir: Path) -> None:
        """Delete all documents and their files. Admin only."""
        try:
            await self.session.execute(delete(DocumentChunk))
            await self.session.execute(delete(Document))
            await self.session.commit()

            for file in upload_dir.glob("*"):
                if file.is_file():
                    file.unlink()

        except Exception as e:
            await self.session.rollback()
            raise FileOperationError(f"Failed to clear documents: {e}")

    async def create_from_file(
        self,
        user: "User",
        filename: str,
        content: bytes,
        upload_dir: Path,
        api_key_id: uuid.UUID | None = None,
    ) -> Document:
        """Create a document record from an uploaded file."""
        safe_filename = sanitize_filename(filename)
        if not safe_filename:
            raise InvalidFilenameError()

        file_ext = Path(safe_filename).suffix.lower()
        if file_ext not in self.ALLOWED_EXTENSIONS:
            raise UnsupportedFileTypeError(file_ext, list(self.ALLOWED_EXTENSIONS))

        unique_filename = f"{uuid.uuid4().hex[:8]}_{safe_filename}"
        safe_path = upload_dir / unique_filename

        try:
            await write_file_async(safe_path, content, mode="wb")

            document = Document(
                id=uuid.uuid4(),
                user_id=user.id,
                api_key_id=api_key_id,
                name=safe_filename,
                file_type=file_ext.lstrip("."),
                file_size=len(content),
                storage_key=str(safe_path),
                status=DocumentStatus.PENDING,
            )
            self.session.add(document)
            await self.session.commit()
            await self.session.refresh(document)

            return document

        except (InvalidFilenameError, UnsupportedFileTypeError):
            raise
        except Exception as e:
            if safe_path.exists():
                safe_path.unlink()
            raise FileOperationError(f"Failed to save file: {e}")

    async def create_from_url(
        self,
        user: "User",
        url: str,
        api_key_id: uuid.UUID | None = None,
    ) -> Document:
        """Create a document record from a URL."""
        if not is_safe_url(url):
            raise InvalidURLError()

        parsed_url = urlparse(url)
        url_filename = f"web_{parsed_url.netloc}_{uuid.uuid4().hex[:8]}.txt"

        document = Document(
            id=uuid.uuid4(),
            user_id=user.id,
            api_key_id=api_key_id,
            name=url_filename,
            file_type="url",
            file_size=0,
            storage_key=url,
            status=DocumentStatus.PENDING,
            extra_metadata={"source_url": url},
        )
        self.session.add(document)
        await self.session.commit()
        await self.session.refresh(document)

        return document

    # Properties (lazy-loaded)

    @property
    def text_splitter(self):
        """Lazy-load text splitter."""
        if self._text_splitter is None:
            from langchain_text_splitters import RecursiveCharacterTextSplitter

            optimal = self.config.get_optimal_settings()
            self._text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=optimal["chunk_size"],
                chunk_overlap=self.config.CHUNK_OVERLAP,
                length_function=len,
                separators=["\n\n", "\n", " ", ""],
            )
        return self._text_splitter

    @property
    def file_loaders(self):
        """Lazy-load file loaders mapping."""
        if self._file_loaders is None:
            from langchain_community.document_loaders import (
                PyPDFLoader,
                TextLoader,
            )

            self._file_loaders = {
                "pdf": PyPDFLoader,
                "txt": TextLoader,
                "text": TextLoader,
                "csv": self._load_csv,
                "md": self._load_markdown,
                "markdown": self._load_markdown,
                "docx": self._load_docx,
                "xlsx": self._load_xlsx,
                "xls": self._load_xlsx,
                "png": self._load_image,
                "jpg": self._load_image,
                "jpeg": self._load_image,
            }
        return self._file_loaders

    async def process_document(
        self,
        document: Document,
        file_path: Path,
        chunk_size: Optional[int] = None,
        progress_callback: Optional[Callable[[str, float, str], None]] = None,
    ) -> dict[str, Any]:
        """
        Process a document and store chunks in pgvector.

        Args:
            document: Document database record
            file_path: Path to the file to process
            chunk_size: Optional custom chunk size
            progress_callback: Optional callback for progress updates (stage, percent, message)

        Returns:
            Processing result dict
        """
        start_time = time.time()

        try:
            # Update document status
            document.status = DocumentStatus.PROCESSING
            await self.session.flush()

            # Stage 1: Load document
            if progress_callback:
                await progress_callback("loading", 10, f"Loading {document.name}")

            loop = asyncio.get_event_loop()
            langchain_docs = await loop.run_in_executor(
                None,
                lambda: self._load_file(str(file_path), document.name),
            )

            # Stage 2: Split into chunks
            if progress_callback:
                await progress_callback("chunking", 30, "Splitting into chunks")

            splitter = self.text_splitter
            if chunk_size:
                from langchain_text_splitters import RecursiveCharacterTextSplitter

                splitter = RecursiveCharacterTextSplitter(
                    chunk_size=chunk_size,
                    chunk_overlap=self.config.CHUNK_OVERLAP,
                    length_function=len,
                    separators=["\n\n", "\n", " ", ""],
                )

            chunks = await loop.run_in_executor(
                None,
                lambda: splitter.split_documents(langchain_docs),
            )

            if not chunks:
                # Create minimal chunk for empty files
                from langchain_core.documents import Document as LCDocument

                chunks = [
                    LCDocument(
                        page_content=f"Empty file: {document.filename}",
                        metadata={"is_empty": True},
                    )
                ]

            # Stage 3: Generate embeddings and store
            if progress_callback:
                await progress_callback("embedding", 50, f"Generating embeddings for {len(chunks)} chunks")

            # Prepare chunks for storage
            chunk_data = []
            for i, chunk in enumerate(chunks):
                chunk_data.append({
                    "content": chunk.page_content,
                    "token_count": len(chunk.page_content) // 4,  # Rough estimate
                    "metadata": {
                        **chunk.metadata,
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "filename": document.name,
                    },
                })

            # Create vector store and add chunks
            vector_store = PgVectorStore(
                session=self.session,
                embedding_provider=self.config.EMBEDDING_PROVIDER,
                embedding_model=self.config.EMBEDDING_MODEL,
            )

            # Generate embeddings in batches to show progress
            batch_size = self.config.BATCH_SIZE
            all_db_chunks = []

            for i in range(0, len(chunk_data), batch_size):
                batch = chunk_data[i : i + batch_size]

                if progress_callback:
                    progress = 50 + (40 * (i + len(batch)) / len(chunk_data))
                    await progress_callback(
                        "embedding",
                        progress,
                        f"Processing chunks {i + 1}-{min(i + len(batch), len(chunk_data))} of {len(chunk_data)}",
                    )

                # Add chunks (this generates embeddings and inserts)
                db_chunks = await vector_store.add_chunks(document.id, batch)
                all_db_chunks.extend(db_chunks)

            # Stage 4: Update document record
            if progress_callback:
                await progress_callback("finalizing", 95, "Finalizing")

            processing_time = time.time() - start_time

            document.status = DocumentStatus.READY
            document.chunk_count = len(all_db_chunks)
            document.extra_metadata = {
                **document.extra_metadata,
                "page_count": len(langchain_docs),
                "processing_time": round(processing_time, 2),
                "embedding_provider": vector_store.provider.value,
            }

            await self.session.flush()

            return {
                "document_id": str(document.id),
                "filename": document.name,
                "pages": len(langchain_docs),
                "chunks": len(all_db_chunks),
                "processing_time": round(processing_time, 2),
                "embedding_provider": vector_store.provider.value,
            }

        except Exception as e:
            document.status = DocumentStatus.FAILED
            document.error_message = str(e)
            await self.session.flush()
            raise

    def _detect_file_type(self, filename: str) -> str:
        """Detect file type from extension."""
        ext = filename.lower().split(".")[-1]
        type_mapping = {
            "pdf": "pdf",
            "txt": "text",
            "csv": "csv",
            "md": "markdown",
            "markdown": "markdown",
            "docx": "docx",
            "pptx": "pptx",
            "xlsx": "xlsx",
            "xls": "xlsx",
            "png": "png",
            "jpg": "jpg",
            "jpeg": "jpg",
        }
        if ext not in type_mapping:
            raise ValueError(f"Unsupported file type: .{ext}")
        return type_mapping[ext]

    def _load_file(self, file_path: str, filename: str) -> list:
        """Load a file using the appropriate loader."""
        file_type = self._detect_file_type(filename)

        if file_type not in self.file_loaders:
            raise ValueError(f"No loader for file type: {file_type}")

        loader = self.file_loaders[file_type]

        if callable(loader) and hasattr(loader, "__name__") and loader.__name__.startswith("_"):
            # Custom loader method
            return loader(file_path)
        else:
            # Standard LangChain loader
            return loader(file_path).load()

    # --- Custom loaders (copied from DocumentProcessor) ---

    def _load_csv(self, file_path: str) -> list:
        """Custom CSV loader."""
        import pandas as pd
        from langchain_core.documents import Document as LCDocument

        documents = []
        try:
            df = pd.read_csv(file_path)
            for idx, row in df.iterrows():
                content_parts = []
                for col, value in row.items():
                    if pd.notna(value):
                        content_parts.append(f"{col}: {value}")
                content = "\n".join(content_parts)
                documents.append(
                    LCDocument(
                        page_content=content,
                        metadata={"source": file_path, "row": idx, "type": "csv"},
                    )
                )
        except Exception:
            from langchain_community.document_loaders import CSVLoader

            loader = CSVLoader(file_path)
            documents = loader.load()
        return documents

    def _load_markdown(self, file_path: str) -> list:
        """Custom Markdown loader."""
        from langchain_core.documents import Document as LCDocument

        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        sections = []
        current_section = []
        for line in content.split("\n"):
            if line.startswith("#") and current_section:
                sections.append("\n".join(current_section))
                current_section = [line]
            else:
                current_section.append(line)
        if current_section:
            sections.append("\n".join(current_section))

        documents = []
        for i, section in enumerate(sections):
            if section.strip():
                documents.append(
                    LCDocument(
                        page_content=section,
                        metadata={"source": file_path, "section": i, "type": "markdown"},
                    )
                )

        if not documents:
            documents.append(
                LCDocument(
                    page_content=content,
                    metadata={"source": file_path, "type": "markdown"},
                )
            )
        return documents

    def _load_docx(self, file_path: str) -> list:
        """Custom Word document loader."""
        from docx import Document as DocxDocument
        from langchain_core.documents import Document as LCDocument

        doc = DocxDocument(file_path)
        content_blocks = []
        current_block = []

        for para in doc.paragraphs:
            if para.text.strip():
                if para.style.name.startswith("Heading"):
                    if current_block:
                        content_blocks.append("\n".join(current_block))
                        current_block = []
                current_block.append(para.text)

        if current_block:
            content_blocks.append("\n".join(current_block))

        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))
            if table_text:
                content_blocks.append("\n".join(table_text))

        documents = []
        for i, block in enumerate(content_blocks):
            if block.strip():
                documents.append(
                    LCDocument(
                        page_content=block,
                        metadata={"source": file_path, "block": i, "type": "docx"},
                    )
                )

        if not documents:
            documents.append(
                LCDocument(
                    page_content="Empty document",
                    metadata={"source": file_path, "type": "docx"},
                )
            )
        return documents

    def _load_xlsx(self, file_path: str) -> list:
        """Custom Excel loader."""
        import openpyxl
        from langchain_core.documents import Document as LCDocument

        documents = []
        wb = openpyxl.load_workbook(file_path, data_only=True)

        for sheet_idx, sheet_name in enumerate(wb.sheetnames):
            sheet = wb[sheet_name]
            if sheet.max_row == 0 or sheet.max_column == 0:
                continue

            sheet_content = [f"Sheet: {sheet_name}", "=" * 50]

            for row in sheet.iter_rows(values_only=True):
                if all(cell is None for cell in row):
                    continue
                row_text = [str(cell).strip() for cell in row if cell is not None]
                if row_text:
                    sheet_content.append(" | ".join(row_text))

            if len(sheet_content) > 2:
                documents.append(
                    LCDocument(
                        page_content="\n".join(sheet_content),
                        metadata={
                            "source": file_path,
                            "sheet_name": sheet_name,
                            "sheet_index": sheet_idx,
                            "type": "xlsx",
                        },
                    )
                )

        if not documents:
            documents.append(
                LCDocument(
                    page_content="Empty Excel file",
                    metadata={"source": file_path, "type": "xlsx"},
                )
            )
        return documents

    def _load_image(self, file_path: str) -> list:
        """Custom image loader with OCR."""
        import pytesseract
        from PIL import Image
        from langchain_core.documents import Document as LCDocument

        try:
            with Image.open(file_path) as image:
                extracted_text = pytesseract.image_to_string(image)

            lines = [line.strip() for line in extracted_text.split("\n") if line.strip()]

            if lines:
                content = f"Image Content (OCR):\n{'-' * 40}\n" + "\n".join(lines)
            else:
                file_ext = file_path.lower().split(".")[-1]
                content = f"Visual image ({file_ext.upper()}): No extractable text"

            return [
                LCDocument(
                    page_content=content,
                    metadata={
                        "source": file_path,
                        "type": file_path.lower().split(".")[-1],
                        "extraction_method": "OCR",
                    },
                )
            ]
        except Exception as e:
            file_ext = file_path.lower().split(".")[-1]
            return [
                LCDocument(
                    page_content=f"Image file ({file_ext.upper()}) - OCR failed: {e}",
                    metadata={"source": file_path, "type": file_ext, "error": str(e)},
                )
            ]


async def get_document_service(session: AsyncSession) -> DocumentService:
    """Factory function for document service."""
    return DocumentService(session=session)
